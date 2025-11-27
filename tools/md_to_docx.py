import sys
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

INPUT = "combined_report.md"
OUTPUT = "Bao_Cao_Do_An_KTPM.docx"

def add_paragraph_with_style(doc, text, bold=False, italic=False, font_name='Times New Roman', size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)


def main():
    doc = Document()
    doc.core_properties.title = 'Đồ án Kiểm Thử Phần Mềm - Báo cáo'

    try:
        with open(INPUT, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Input file '{INPUT}' not found.")
        sys.exit(1)

    in_code = False
    code_buffer = []
    para_buffer = []

    def flush_para():
        nonlocal para_buffer
        if not para_buffer:
            return
        text = ' '.join([ln.strip() for ln in para_buffer])
        add_paragraph_with_style(doc, text)
        para_buffer = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        # add code block as a paragraph with monospace font
        p = doc.add_paragraph()
        for ln in code_buffer:
            run = p.add_run(ln.rstrip('\n') + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        code_buffer = []

    for raw in lines:
        line = raw.rstrip('\n')
        if line.strip().startswith('```'):
            if in_code:
                # end code
                in_code = False
                flush_code()
            else:
                # start code
                in_code = True
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line + '\n')
            continue

        if line.strip() == '':
            flush_para()
            continue

        if line.startswith('# '):
            flush_para()
            heading = line[2:].strip()
            doc.add_heading(heading, level=1)
            continue
        if line.startswith('## '):
            flush_para()
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            continue
        if line.startswith('### '):
            flush_para()
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
            continue

        # simple table or list line - just append to paragraph buffer
        para_buffer.append(line)

    flush_para()
    flush_code()

    try:
        doc.save(OUTPUT)
        print(f"Saved {OUTPUT}")
    except Exception as e:
        print("Error saving DOCX:", e)
        sys.exit(1)

if __name__ == '__main__':
    main()
